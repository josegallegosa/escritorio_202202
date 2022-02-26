# -*- coding: utf-8 -*-
from odoo import models, fields, api
from openpyxl import Workbook, load_workbook
from openpyxl.styles import NamedStyle, PatternFill, Border, Side, Alignment, Protection, Font
import os
from odoo.exceptions import ValidationError
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Inches      
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Cm
from docx.enum.style import (WD_BUILTIN_STYLE, WD_STYLE, WD_STYLE_TYPE)
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_BREAK
import datetime


def font_size_docx(rows,pt):
    for row in rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= pt


def make_column_bold(column,*rows):
    cont = 1
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if cont == column:
                    cont+=1
                    for run in paragraph.runs:
                        run.font.bold = True

def make_row_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True




class HRPayrollReportSumaryButtonXML(models.Model):
    _inherit = 'hr.payslip'

    date_today = fields.Datetime.now().strftime("%Y-%m-%d")#llamarlo directo 
    time_today = fields.Datetime.now().strftime("%H:%M:%S")#llamarlo directo 
    document_number_RUC = fields.Char(string='Número de Documento')  #este ya está en company.id 
    order_number = fields.Char(string='Número de Orden') # se llama number en planilla 

    
    def print_docx(self):
        RUTA_BASE = os.path.dirname(os.path.abspath(__file__))
        document = Document()
        section = document.sections[0]
        header = section.header
        header.is_linked_to_previous = True
        
        paragraph1 = document.add_paragraph()
        title1 = ('LIQUIDACIÓN DE BENEFICIOS SOCIALES')
        p_title1 = paragraph1.add_run(title1)
        font = p_title1.font
        font.size = Pt(11)
        font.color.rgb = RGBColor(000, 000, 000)
        p_title1.bold = True
        paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraph1.paragraph_format.line_spacing = 0.8

        paragraph2 = document.add_paragraph()
        title2 = ('D.L. N° 650, D. L. N° 857 y D.S. N° 001-97-TR')
        p_title2 = paragraph2.add_run(title2)
        font = p_title1.font
        font.size = Pt(9)
        font.color.rgb = RGBColor(000, 000, 000)
        p_title2.bold = False
        paragraph2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraph2.paragraph_format.line_spacing = 0.8

        records = (
            ('Código:','87654321'),
            ('DNI:','87654321'),
            ('Apellidos y Nombres: ','NOMBRE DEL TRABAJADOR'),
            ('Cargo:','Jefe de Recursos Humanos'),
            ('Básico:','S/7,500.00'),
            ('Condición:','Empleado'),
            ('Pensión:','Prima'),
            ('Fecha de Ingreso:','20/11/2020'),
            ('Fecha de Cese:','31/01/2021'),
            ('Tiempo de Servicios:','2 Meses; 12 Días'),
            ('Motivo de Cese:','NO PASO PERIODO DE PRUEBA'),
            ('Régimen Laboral:','PRIVADO GENERAL - D. LEG. N.° 728'),
        )

        table = document.add_table(rows=1, cols=2, style=None)
        table.autofit = False 
        table.allow_autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for col1, col2 in records:
            row_cells = table.add_row().cells
            row_cells[0].text = col1
            row_cells[1].text = col2
        
        font_size_docx(table.rows, Pt(8))

        for x in range(13):
            make_column_bold(1,table.rows[x])

        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        """
        eb_1 = document.add_paragraph('')
        eb_1.paragraph_format.line_spacing = 0
        """
        paragraph3 = document.add_paragraph()
        title3 = ('Conceptos Remunerativos')
        p_title3 = paragraph3.add_run(title3)
        font = p_title3.font
        font.size = Pt(10)
        font.color.rgb = RGBColor(000, 000, 000)
        p_title3.bold = True
        p_title3.underline = True
        paragraph3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        

        #table2

        records2 = (
            ('1/6 de Gratificación','210.92','0.00','0.00'),
            ('Asignación Familiar','93.00','93.00','93.00'),
            ('Básico','7,500.00','7,500.00','7,500.00'),
            ('TOTAL REMUNERACION COMPUTABLE','7,803.92','7,593.00','7,593.00'),
        )

        table2 = document.add_table(rows=1, cols=4, style=None)
        table2.autofit = False 
        table2.allow_autofit  = False
        table2.columns[0].width = Cm(4.19)
        table2.columns[1].width = Cm(2.14)
        table2.columns[2].width = Cm(2.14)
        table2.columns[3].width = Cm(2.14)
        table2.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].paragraphs[0].text = 'C.T.S'
        hdr_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT 
        hdr_cells[2].paragraphs[0].text = 'GRATIFICACIÓN'
        hdr_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        hdr_cells[3].paragraphs[0].text = 'VACACIONES'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for col1, col2, col3, col4 in records2:
            row_cells = table2.add_row().cells
            row_cells[0].text = col1
            row_cells[1].paragraphs[0].text = col2
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT 
            row_cells[2].paragraphs[0].text = col3
            row_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[3].paragraphs[0].text = col4
            row_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        
        font_size_docx(table2.rows, Pt(8))

        make_row_bold(table2.rows[0],table2.rows[4])

        for row in table2.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        
        eb_2 = document.add_paragraph('')
        eb_2.paragraph_format.line_spacing = 0
        
        #table3

        records3 = (
            ('Periodo por Liquidar: 0 Años 2 Meses 12 Días','','',''),
            ('Del 20 de Noviembre del 2020 al 31 de Enero del 2021','','',''),
            ('-7803.92 / 12 * 2 meses laborados','=','S/1,300.65',''),
            ('-7803.92 / 12 / 30 * 12 días laborados','=','S/260.13',''),
        )

        table3 = document.add_table(rows=1, cols=4, style=None)
        table3.autofit = False 
        table3.allow_autofit = False
        table3.columns[0].width = Cm(4.19)
        table3.columns[1].width = Cm(2.14)
        table3.columns[2].width = Cm(2.14)
        table3.columns[3].width = Cm(2.14)
        table3.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table3.rows[0].cells
        hdr_cells[0].text = '1. COMPENSACIÓN POR TIEMPO DE SERVICIO'
        hdr_cells[1].text = ''
        hdr_cells[2].text = ''
        hdr_cells[3].paragraphs[0].text = 'S/ 1,560.78'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for col1, col2, col3, col4 in records3:
            row_cells = table3.add_row().cells
            row_cells[0].text = col1
            row_cells[1].paragraphs[0].text = col2
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[2].text = col3
            row_cells[3].text = col4
        
        font_size_docx(table3.rows, Pt(8))

        make_row_bold(table3.rows[0])

        for row in table3.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        """
        eb_3 = document.add_paragraph('')
        eb_3.paragraph_format.line_spacing = 0
        """
        #table4

        records4 = (
            ('Periodo por Liquidar: 0 Años 1 Meses 0 Días','','',''),
            ('Del 01 de Enero del 2021 al 31 de Enero del 2021','','',''),
            ('-7593.00 / 6 * 1 meses laborados','=','S/1,265.50',''),
            ('-Bonificación Ext. 6.75% ','=','S/85.42',''),
        )

        table4= document.add_table(rows=1, cols=4, style=None)
        table4.autofit = False 
        table4.allow_autofit = False
        table4.columns[0].width = Cm(4.19)
        table4.columns[1].width = Cm(2.14)
        table4.columns[2].width = Cm(2.14)
        table4.columns[3].width = Cm(2.14)
        table4.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table4.rows[0].cells
        hdr_cells[0].text = '2. GRATIFICACIÓN TRUNCA'
        hdr_cells[1].text = ''
        hdr_cells[2].text = ''
        hdr_cells[3].paragraphs[0].text = 'S/ 1,350.92'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for col1, col2, col3, col4 in records4:
            row_cells = table4.add_row().cells
            row_cells[0].text = col1
            row_cells[1].paragraphs[0].text = col2
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[2].text = col3
            row_cells[3].text = col4
        
        font_size_docx(table4.rows, Pt(8))

        make_row_bold(table4.rows[0])

        for row in table4.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        """
        eb_4 = document.add_paragraph('')
        eb_4.paragraph_format.line_spacing = 0
        """
        #table5

        records5 = (
            ('Periodo por Liquidar: 0 Años 2 Meses 12 Días','','',''),
            ('Del 20 de Noviembre del 2020 al 31 de Enero del 2021','','',''),
            ('-7593.00 / 12 * 2 meses laborados','=','S/1,265.50',''),
            ('-7593.00 / 12 / 30 * 12 días laborados','=','S/253.10',''),
        )

        table5= document.add_table(rows=1, cols=4, style=None)
        table5.autofit = False 
        table5.allow_autofit = False
        table5.columns[0].width = Cm(4.19)
        table5.columns[1].width = Cm(2.14)
        table5.columns[2].width = Cm(2.14)
        table5.columns[3].width = Cm(2.14)
        table5.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table5.rows[0].cells
        hdr_cells[0].text = '3. VACACIONES TRUNCAS'
        hdr_cells[1].text = ''
        hdr_cells[2].text = ''
        hdr_cells[3].paragraphs[0].text = 'S/ 1,518.60'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for col1, col2, col3, col4 in records5:
            row_cells = table5.add_row().cells
            row_cells[0].text = col1
            row_cells[1].paragraphs[0].text = col2
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[2].text = col3
            row_cells[3].text = col4
        
        font_size_docx(table5.rows, Pt(8))

        make_row_bold(table5.rows[0])

        for row in table5.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        """
        eb_5 = document.add_paragraph('')
        eb_5.paragraph_format.line_spacing = 0
        """
        #table6

        records6 = (
            ('Asignación Familiar','=','S/93.00',''),
            ('-7593.00 / 12 / 30 * 12 días laborados','=','S/7,500.00',''),
        )

        table6= document.add_table(rows=1, cols=4, style=None)
        table6.autofit = False 
        table6.allow_autofit = False
        table6.columns[0].width = Cm(4.19)
        table6.columns[1].width = Cm(2.14)
        table6.columns[2].width = Cm(2.14)
        table6.columns[3].width = Cm(2.14)
        table6.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table6.rows[0].cells
        hdr_cells[0].text = '4. REMUNERACIONES'
        hdr_cells[1].text = ''
        hdr_cells[2].text = ''
        hdr_cells[3].paragraphs[0].text = 'S/ 7,593.00'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for col1, col2, col3, col4 in records6:
            row_cells = table6.add_row().cells
            row_cells[0].text = col1
            row_cells[1].paragraphs[0].text = col2
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[2].text = col3
            row_cells[3].text = col4
        
        font_size_docx(table6.rows, Pt(8))

        make_row_bold(table6.rows[0])

        for row in table6.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        """
        eb_6 = document.add_paragraph('')
        eb_6.paragraph_format.line_spacing = 0
        """
        #table7

        records7 = (
            ('AFP. Aportación Obligator '+'(10.00%)','=','S/93.00',''),
            ('AFP. Seguro de Vida '+'(1.74%)','=','S/7,500.00',''),
            ('AFP. Comisión sobre la RA '+'(1.60%)','=','S/7,500.00',''),
        )

        table7= document.add_table(rows=1, cols=4, style=None)
        table7.autofit = False 
        table7.allow_autofit = False
        table7.columns[0].width = Cm(4.19)
        table7.columns[1].width = Cm(2.14)
        table7.columns[2].width = Cm(2.14)
        table7.columns[3].width = Cm(2.14)
        table7.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table7.rows[0].cells
        hdr_cells[0].text = '5. DEDUCCIONES'
        hdr_cells[1].text = ''
        hdr_cells[2].text = ''
        hdr_cells[3].paragraphs[0].text = 'S/ (1,215.49)'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for col1, col2, col3, col4 in records7:
            row_cells = table7.add_row().cells
            row_cells[0].text = col1
            row_cells[1].paragraphs[0].text = col2
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[2].text = col3
            row_cells[3].text = col4
        
        font_size_docx(table7.rows, Pt(8))

        make_row_bold(table7.rows[0])

        for row in table7.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        """
        eb_7 = document.add_paragraph('')
        eb_7.paragraph_format.line_spacing = 0
        """
        #table6

        records8 = (
            ('EPS','=','S/93.00',''),
            ('<Adelanto de Quincena>','=','S/2,657.00',''),
        )

        table8= document.add_table(rows=1, cols=4, style=None)
        table8.autofit = False 
        table8.allow_autofit = False
        table8.columns[0].width = Cm(4.19)
        table8.columns[1].width = Cm(2.14)
        table8.columns[2].width = Cm(2.14)
        table8.columns[3].width = Cm(2.14)
        table8.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table8.rows[0].cells
        hdr_cells[0].text = '6. OTROS Descuentos'
        hdr_cells[1].text = ''
        hdr_cells[2].text = ''
        hdr_cells[3].paragraphs[0].text = 'S/ (3,003.13)'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for col1, col2, col3, col4 in records8:
            row_cells = table8.add_row().cells
            row_cells[0].text = col1
            row_cells[1].paragraphs[0].text = col2
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[2].text = col3
            row_cells[3].text = col4
        
        font_size_docx(table8.rows, Pt(8))

        make_row_bold(table8.rows[0])

        for row in table8.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        """
        eb_8 = document.add_paragraph('')
        eb_8.paragraph_format.line_spacing = 0
        """
        #table6

        records9 = (
            ('EPS APORT '+'(2.25%)','=','S/205.01',''),
            ('EsSalud '+'(6.75%)','=','S/615.03',''),
        )

        table9= document.add_table(rows=1, cols=4, style=None)
        table9.autofit = False 
        table9.allow_autofit = False
        table9.columns[0].width = Cm(4.19)
        table9.columns[1].width = Cm(2.14)
        table9.columns[2].width = Cm(2.14)
        table9.columns[3].width = Cm(2.14)
        table9.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table9.rows[0].cells
        hdr_cells[0].text = '7. CONTRIBUCIONES'
        hdr_cells[1].text = ''
        hdr_cells[2].text = ''
        hdr_cells[3].paragraphs[0].text = 'S/ (820.04)'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for col1, col2, col3, col4 in records9:
            row_cells = table9.add_row().cells
            row_cells[0].text = col1
            row_cells[1].paragraphs[0].text = col2
            row_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[2].text = col3
            row_cells[3].text = col4
        
        font_size_docx(table9.rows, Pt(8))

        make_row_bold(table9.rows[0])

        for row in table9.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        """
        eb_9 = document.add_paragraph('')
        eb_9.paragraph_format.line_spacing = 0
        """
        #table6
     

        table10= document.add_table(rows=1, cols=4, style=None)
        table10.autofit = False 
        table10.allow_autofit = False
        table10.columns[0].width = Cm(4.19)
        table10.columns[1].width = Cm(2.14)
        table10.columns[2].width = Cm(2.14)
        table10.columns[3].width = Cm(2.14)
        table10.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table10.rows[0].cells
        hdr_cells[0].merge(hdr_cells[2]).paragraphs[0].text = 'NETO A PAGAR SOLES'
        hdr_cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].paragraphs[0].text = 'S/ 7,804.68'
        hdr_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        

        
        font_size_docx(table10.rows, Pt(8))

        make_row_bold(table10.rows[0])

        for row in table10.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)
        
        
        eb_10 = document.add_paragraph('')
        eb_10.paragraph_format.line_spacing = 0
        
        paragraph4 = document.add_paragraph()
        title4 = ('Recibí de la empresa COOP. DE SERV. EDUC. ABRAHAM LINCOLN LTDA, la suma de ')
        title4_2 = ('S/ 7,804.68')
        title4_3 = (' (Siete Mil Ochocientos Cuatro con 68/100 Soles), correspondiente a mis beneficios sociales conforme a ley, firmo en señal de conformidad.')
        p_title4 = paragraph4.add_run(title4)
        p_title4_2 = paragraph4.add_run(title4_2)
        p_title4_3 = paragraph4.add_run(title4_3)
        font = p_title4.font
        font_2 = p_title4_2.font
        font_3 = p_title4_3.font
        font.size = Pt(8)
        font_2.size = Pt(8)
        font_3.size = Pt(8)
        font.color.rgb = RGBColor(000, 000, 000)
        paragraph4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        paragraph4.paragraph_format.line_spacing = 0.8

        """
        eb_10 = document.add_paragraph('')
        eb_10.paragraph_format.line_spacing = 0
        """

        paragraph5 = document.add_paragraph()
        title5 = ('LIMA, 21 de Enero de 2021')
        p_title5 = paragraph5.add_run(title5)
        font = p_title5.font
        font.size = Pt(8)
        font.color.rgb = RGBColor(000, 000, 000)
        paragraph5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        paragraph5.paragraph_format.line_spacing = 0.8

        eb_11 = document.add_paragraph('')
        eb_11.paragraph_format.line_spacing = 0

        records11 = (
            ('NOMBRE: '+'APODERADO','','NOMBRE: '+'TRABAJADOR',''),
            ('DNI: '+'12345678','','DNI Nro: '+'87654321',''),
        )

        table11= document.add_table(rows=1, cols=4, style=None)
        table11.autofit = False 
        table11.allow_autofit = False
        table11.columns[0].width = Cm(4.19)
        table11.columns[1].width = Cm(2.14)
        table11.columns[2].width = Cm(2.14)
        table11.columns[3].width = Cm(2.14)
        table11.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table11.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1]).paragraphs[0].text = '____________________________________'
        hdr_cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].merge(hdr_cells[3]).paragraphs[0].text = '____________________________________'
        hdr_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        for col1, col2, col3, col4 in records11:
            row_cells = table11.add_row().cells
            row_cells[0].merge(row_cells[1]).paragraphs[0].text = col1
            row_cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].merge(row_cells[3]).paragraphs[0].text = col3
            row_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        
        font_size_docx(table11.rows, Pt(7))


        for row in table11.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Pt(10)


        name = 'prueba.docx'

        document.save(RUTA_BASE+"/../static/src/archivos/"+name)
        

        return {
        'name': "Resumen de Nómina",
        'type': 'ir.actions.act_url',
        'url': self.env['ir.config_parameter'].get_param('web.base.url')+'/rapitech_hr_report_quanam/static/src/archivos/'+name,
        'target': 'new',
        }